import os, struct, io, zipfile
from datetime import datetime
from xml.sax.saxutils import escape as xml_escape

try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False

MESES = ["", "enero", "febrero", "marzo", "abril", "mayo", "junio",
         "julio", "agosto", "setiembre", "octubre", "noviembre", "diciembre"]

CM2EMU = 360000
INCH2EMU = 914400
PT2HPT = 2  # half-points


def _fmt_fecha(fecha_str):
    try:
        d = datetime.strptime(fecha_str, "%Y-%m-%d")
        return d.strftime("%d/%m/%Y")
    except (ValueError, TypeError):
        return fecha_str


def _fmt_fecha_larga(fecha_str):
    try:
        d = datetime.strptime(fecha_str, "%Y-%m-%d")
        return f"{d.day:02d} de {MESES[d.month]} del {d.year}"
    except (ValueError, TypeError):
        return fecha_str


def _img_dims(path):
    try:
        with open(path, 'rb') as f:
            h = f.read(8)
            if h[:8] == b'\x89PNG\r\n\x1a\n':
                f.read(4)
                f.read(4)
                w, h = struct.unpack('>II', f.read(8))
                return w, h
            if h[:2] != b'\xff\xd8':
                return None
            f.seek(2)
            while True:
                b = f.read(1)
                while b and b[0] == 0xFF:
                    b = f.read(1)
                if not b:
                    return None
                marker = b[0]
                if marker in (0xC0, 0xC2):
                    f.read(3)
                    h_b, w_b = struct.unpack('>HH', f.read(4))
                    return w_b, h_b
                if marker in (0xD9, 0xDA):
                    return None
                if marker in (0x01,) or (0xD0 <= marker <= 0xD7) or marker == 0xD8:
                    continue
                ln_data = f.read(2)
                if len(ln_data) < 2:
                    return None
                ln = struct.unpack('>H', ln_data)[0]
                if ln < 2:
                    return None
                f.seek(ln - 2, 1)
    except Exception:
        return None


def _emu(val, unit='cm'):
    if unit == 'cm':
        return int(val * CM2EMU)
    return int(val * INCH2EMU)


def _make_rels(*items):
    lines = ['<?xml version="1.0" encoding="UTF-8" standalone="yes"?>']
    lines.append('<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">')
    for rid, typ, target in items:
        lines.append(f'<Relationship Id="{rid}" Type="{typ}" Target="{target}"/>')
    lines.append('</Relationships>')
    return '\n'.join(lines)


def _mkpara(text, bold=False, size=20, color="000000", font="Calibri", align=None, underline=False, spacing_after=100):
    parts = []
    parts.append('<w:p>')
    parts.append('<w:pPr>')
    parts.append(f'<w:spacing w:after="{spacing_after}" w:line="240" w:lineRule="auto"/>')
    if align:
        parts.append(f'<w:jc w:val="{align}"/>')
    parts.append('</w:pPr>')
    parts.append('<w:r>')
    parts.append('<w:rPr>')
    if bold:
        parts.append('<w:b/>')
    if underline:
        parts.append('<w:u w:val="single"/>')
    parts.append(f'<w:sz w:val="{size}"/>')
    parts.append(f'<w:rFonts w:ascii="{font}" w:hAnsi="{font}"/>')
    parts.append(f'<w:color w:val="{color}"/>')
    parts.append('</w:rPr>')
    parts.append(f'<w:t xml:space="preserve">{xml_escape(text)}</w:t>')
    parts.append('</w:r>')
    parts.append('</w:p>')
    return '\n'.join(parts)


def _mkfield(label, value_lines):
    parts = []
    parts.append('<w:p>')
    parts.append('<w:pPr><w:spacing w:after="60" w:before="0"/></w:pPr>')
    parts.append('<w:r>')
    parts.append('<w:rPr><w:b/><w:sz w:val="20"/><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/></w:rPr>')
    parts.append(f'<w:t xml:space="preserve">{xml_escape(label)}&#9;: </w:t>')
    parts.append('</w:r>')
    for i, line in enumerate(value_lines):
        if i > 0:
            parts.append('<w:r><w:br/></w:r>')
        parts.append('<w:r>')
        parts.append('<w:rPr><w:sz w:val="20"/><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/></w:rPr>')
        parts.append(f'<w:t xml:space="preserve">{xml_escape(line)}</w:t>')
        parts.append('</w:r>')
    parts.append('</w:p>')
    return '\n'.join(parts)


def _mksection_heading(text):
    return _mkpara(text, bold=True, size=20, color="000000")


def _mktable(hdrs, rows, col_widths_twips=None):
    parts = []
    parts.append('<w:tbl>')
    parts.append('<w:tblPr>')
    parts.append('<w:tblStyle w:val="LightShading-Accent1"/>')
    parts.append('<w:tblW w:w="5000" w:type="pct"/>')
    parts.append('<w:jc w:val="center"/>')
    parts.append('</w:tblPr>')
    if col_widths_twips:
        parts.append('<w:tblGrid>')
        for cw in col_widths_twips:
            parts.append(f'<w:gridCol w:w="{cw}"/>')
        parts.append('</w:tblGrid>')

    def _cell(text, bold=False, sz=16):
        c = ['<w:tc>']
        c.append('<w:tcPr>')
        c.append('<w:vAlign w:val="center"/>')
        if col_widths_twips:
            pass
        c.append('</w:tcPr>')
        c.append('<w:p>')
        c.append('<w:pPr><w:jc w:val="center"/></w:pPr>')
        c.append('<w:r>')
        c.append('<w:rPr>')
        if bold:
            c.append('<w:b/>')
        c.append(f'<w:sz w:val="{sz}"/>')
        c.append('<w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/>')
        c.append('<w:color w:val="000000"/>')
        c.append('</w:rPr>')
        c.append(f'<w:t xml:space="preserve">{xml_escape(str(text))}</w:t>')
        c.append('</w:r>')
        c.append('</w:p>')
        c.append('</w:tc>')
        return '\n'.join(c)

    parts.append('<w:tr>')
    for h in hdrs:
        parts.append(_cell(h, bold=True, sz=16))
    parts.append('</w:tr>')

    for row in rows:
        parts.append('<w:tr>')
        for v in row:
            parts.append(_cell(v, bold=False, sz=16))
        parts.append('</w:tr>')

    parts.append('</w:tbl>')
    return '\n'.join(parts)


def _mkimage_para(r_id, width_emu, height_emu, descr=""):
    return f'''<w:p>
  <w:pPr><w:jc w:val="center"/></w:pPr>
  <w:r>
    <w:rPr/>
    <w:drawing>
      <wp:inline distT="0" distB="0" distL="0" distR="0">
        <wp:extent cx="{width_emu}" cy="{height_emu}"/>
        <wp:effectExtent l="0" t="0" b="0" r="0"/>
        <wp:docPr id="0" name="Picture" descr="{xml_escape(descr)}"/>
        <a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
          <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
            <pic:pic xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">
              <pic:nvPicPr>
                <pic:cNvPr id="0" name="Picture"/>
                <pic:cNvPicPr/>
              </pic:nvPicPr>
              <pic:blipFill>
                <a:blip r:embed="{r_id}"/>
                <a:stretch><a:fillRect/></a:stretch>
              </pic:blipFill>
              <pic:spPr>
                <a:xfrm><a:off x="0" y="0"/><a:ext cx="{width_emu}" cy="{height_emu}"/></a:xfrm>
                <a:prstGeom prst="rect"/>
              </pic:spPr>
            </pic:pic>
          </a:graphicData>
        </a:graphic>
      </wp:inline>
    </w:drawing>
  </w:r>
</w:p>'''


def _build_docx_stdlib(doc_title, fecha_label, sections_data,
                       residente_nombre, residente_cargo,
                       responsable_nombre, responsable_cargo,
                       asunto, proyecto_nombre, cui, clima, observaciones,
                       show_fecha_col):
    buf = io.BytesIO()
    zf = zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED)

    def _add(name, content):
        if isinstance(content, str):
            zf.writestr(name, content.encode('utf-8'))
        else:
            zf.writestr(name, content)

    img_path = os.path.join(os.path.dirname(__file__), "ENCABEZADO.png")
    has_header_img = os.path.exists(img_path)

    body_parts = []
    body_parts.append('__SECTPR__')

    body_parts.append(_mkpara(doc_title, bold=True, size=24, align="center", underline=True,
                               color="000000", spacing_after=200))

    if residente_nombre:
        lines = [residente_nombre]
        if residente_cargo:
            lines.append(residente_cargo)
        body_parts.append(_mkfield("A", lines))

    if responsable_nombre:
        lines = [responsable_nombre]
        if responsable_cargo:
            lines.append(responsable_cargo)
        body_parts.append(_mkfield("DE", lines))

    if asunto:
        body_parts.append(_mkfield("ASUNTO", [asunto]))

    body_parts.append(_mkfield("FECHA", [f"Santo Domingo de Acobamba, {fecha_label}"]))
    body_parts.append(_mkpara("", spacing_after=200))

    body_parts.append(_mksection_heading("DATOS GENERALES"))
    body_parts.append(_mkfield("NOMBRE DEL PROYECTO", [proyecto_nombre]))
    body_parts.append(_mkfield("CUI", [cui]))
    if clima:
        body_parts.append(_mkfield("CLIMA", [clima]))
    if observaciones:
        body_parts.append(_mkfield("OBSERVACIONES", [observaciones]))

    photo_rels = []
    photo_idx = 0
    hdr_ref = None

    for section_data in sections_data:
        s_type = section_data.get("type")

        if s_type == "partidas":
            body_parts.append(_mkpara("", spacing_after=100))
            body_parts.append(_mksection_heading(section_data["title"]))
            avances = section_data["data"]
            if avances:
                hdrs = ["Sector", "Código", "Partida", "Ejecutado por", "Operarios", "Oficiales", "Peones", "Horas", "Cant. Ejec."]
                if show_fecha_col:
                    hdrs = ["Fecha"] + hdrs
                rows = []
                for a in avances:
                    vals = []
                    if show_fecha_col:
                        vals.append(a.get("fecha", ""))
                    vals += [
                        a.get("sector_nombre") or "",
                        a.get("partida_codigo") or "",
                        a["partida_nombre"],
                        a.get("subcontrata_nombre") or "Empresa",
                        str(a.get("num_operarios", 0)),
                        str(a.get("num_oficiales", 0)),
                        str(a.get("num_peones", 0)),
                        f'{a.get("horas_trabajadas", 0):.1f}',
                        f'{a.get("cantidad_ejecutada", 0):.2f}',
                    ]
                    rows.append(vals)
                col_widths = None
                body_parts.append(_mktable(hdrs, rows, col_widths))
                body_parts.append(_mkpara("", spacing_after=100))

        elif s_type == "materiales":
            body_parts.append(_mkpara("", spacing_after=100))
            body_parts.append(_mksection_heading(section_data["title"]))
            materiales = section_data["data"]
            if materiales:
                hdrs = ["Material", "Cantidad", "Unidad", "Fecha"]
                rows = [[m["material"], str(m["cantidad"]), m["unidad"], m["fecha"]] for m in materiales]
                body_parts.append(_mktable(hdrs, rows))
                body_parts.append(_mkpara("", spacing_after=100))

        elif s_type == "notas":
            body_parts.append(_mkpara("", spacing_after=100))
            body_parts.append(_mksection_heading(section_data["title"]))
            for nota in section_data["data"]:
                body_parts.append(_mkpara(f"- {nota['nota']}", size=20, spacing_after=40))

        elif s_type == "fotos":
            body_parts.append(_mkpara("", spacing_after=100))
            body_parts.append(_mksection_heading(section_data["title"]))
            for idx_foto, foto in enumerate(section_data["data"], start=1):
                body_parts.append(_mkpara(f"FOTOGRAFIA N° {idx_foto}", bold=True, size=20,
                                           align="center", spacing_after=40))
                if os.path.exists(foto["ruta"]):
                    dims = _img_dims(foto["ruta"])
                    if dims:
                        fw, fh = dims
                        target_w = int(4.5 * INCH2EMU)
                        target_h = int(target_w * fh / fw)
                        r_id = f"rPhoto{photo_idx}"
                        ext = os.path.splitext(foto["ruta"])[1].lower()
                        fname = f"photo_{photo_idx}{ext}"
                        photo_rels.append((r_id, fname))
                        body_parts.append(_mkimage_para(r_id, target_w, target_h, f"Foto {idx_foto}"))
                        with open(foto["ruta"], 'rb') as fimg:
                            _add(f"word/media/{fname}", fimg.read())
                        photo_idx += 1
                    else:
                        body_parts.append(_mkpara("[No se pudo insertar la imagen]", size=18, align="center"))
                if foto.get("descripcion"):
                    body_parts.append(_mkpara(f"Descripcion: {foto['descripcion']}", size=20, align="center"))

    body_parts.append(_mkpara("", spacing_after=200))
    body_parts.append(_mkpara("", spacing_after=200))
    body_parts.append(_mkpara("_________________________________", size=20, align="center", spacing_after=40))
    body_parts.append(_mkpara(responsable_nombre, bold=True, size=20, align="center", spacing_after=20))
    body_parts.append(_mkpara(responsable_cargo, size=20, align="center"))

    header_rels = []
    if has_header_img:
        with open(img_path, 'rb') as f:
            _add('word/media/encabezado.png', f.read())
        dims = _img_dims(img_path)
        if dims:
            hw, hh = dims
            target_hw = int(15 * CM2EMU)
            target_hh = int(target_hw * hh / hw)
            header_xml = '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            header_xml += ('<w:hdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
                           ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"'
                           ' xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"'
                           ' xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
                           ' xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">')
            header_xml += _mkimage_para("rHeaderImg", target_hw, target_hh, "encabezado")
            header_xml += '</w:hdr>'
            _add('word/header1.xml', header_xml)
            header_rels.append(('rHeaderImg', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image', 'media/encabezado.png'))
            hdr_ref = ('rIdHdr', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/header', 'header1.xml')

    sectPr = ('<w:sectPr>'
              '<w:pgSz w:w="11906" w:h="16838"/>'
              '<w:pgMar w:left="1417" w:right="1417" w:top="1417" w:bottom="1417"/>')
    if hdr_ref:
        sectPr += f'<w:headerReference w:type="default" r:id="{hdr_ref[0]}"/>'
    sectPr += '</w:sectPr>'
    body_parts[0] = sectPr
    doc_xml = '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
    doc_xml += ('<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
                ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"'
                ' xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"'
                ' xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
                ' xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">')
    doc_xml += '<w:body>'
    doc_xml += '\n'.join(body_parts)
    doc_xml += '</w:body></w:document>'
    _add('word/document.xml', doc_xml)

    rels_items = [
        ('rId1', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles', 'styles.xml'),
    ]
    if hdr_ref:
        rels_items.append(('rIdHdr', hdr_ref[1], hdr_ref[2]))
    for r_id, fname in photo_rels:
        rels_items.append((r_id, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image', f'media/{fname}'))
    _add('word/_rels/document.xml.rels', _make_rels(*rels_items))

    if has_header_img:
        _add('word/_rels/header1.xml.rels', _make_rels(*header_rels))

    ct_xml = '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
    ct_xml += '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
    ct_xml += '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
    ct_xml += '<Default Extension="xml" ContentType="application/xml"/>'
    ct_xml += '<Default Extension="png" ContentType="image/png"/>'
    ct_xml += '<Default Extension="jpg" ContentType="image/jpeg"/>'
    ct_xml += '<Default Extension="jpeg" ContentType="image/jpeg"/>'
    ct_xml += '<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
    ct_xml += '<Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>'
    if has_header_img:
        ct_xml += '<Override PartName="/word/header1.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.header+xml"/>'
    ct_xml += '</Types>'
    _add('[Content_Types].xml', ct_xml)

    _add('_rels/.rels', _make_rels(
        ('rId1', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument', 'word/document.xml')
    ))

    styles_xml = '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
    styles_xml += '<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    styles_xml += '<w:style w:type="table" w:styleId="LightShading-Accent1">'
    styles_xml += '<w:name w:val="Light Shading Accent 1"/>'
    styles_xml += '<w:style w:type="paragraph" w:styleId="Normal"><w:name w:val="Normal"/></w:style>'
    styles_xml += '</w:styles>'
    _add('word/styles.xml', styles_xml)

    app_xml = '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
    app_xml += '<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties">'
    app_xml += '<Application>ControlObra</Application></Properties>'
    _add('docProps/app.xml', app_xml)

    core_xml = '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
    core_xml += '<cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties"'
    core_xml += ' xmlns:dc="http://purl.org/dc/elements/1.1/">'
    core_xml += '<dc:title>ControlObra</dc:title></cp:coreProperties>'
    _add('docProps/core.xml', core_xml)

    zf.close()
    buf.seek(0)
    return buf


def _set_col_widths(table, widths):
    if not _HAS_DOCX:
        return
    for i, w in enumerate(widths):
        for row in table.rows:
            row.cells[i].width = Cm(w)


def _build_doc(doc_title, fecha_label, sections_data,
               residente_nombre, residente_cargo,
               responsable_nombre, responsable_cargo,
               asunto, proyecto_nombre, cui, clima, observaciones,
               show_fecha_col):
    if not _HAS_DOCX:
        return _build_docx_stdlib(doc_title, fecha_label, sections_data,
                                   residente_nombre, residente_cargo,
                                   responsable_nombre, responsable_cargo,
                                   asunto, proyecto_nombre, cui, clima, observaciones,
                                   show_fecha_col)
    doc = Document()

    section = doc.sections[0]
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    header = section.header
    header.is_linked_to_previous = False
    p_header = header.paragraphs[0]
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_path = os.path.join(os.path.dirname(__file__), "ENCABEZADO.png")
    if os.path.exists(img_path):
        run = p_header.add_run()
        run.add_picture(img_path, width=Cm(15))

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    def add_heading(text, level=1, size=10):
        p = doc.add_heading(text, level=level)
        for run in p.runs:
            run.font.size = Pt(size)
            run.font.name = "Calibri"
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        return p

    def add_para(text, bold=False, size=10, space_after=Pt(4)):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = "Calibri"
        p.paragraph_format.space_after = space_after
        p.paragraph_format.space_before = Pt(0)
        return p

    def add_field(label, value_lines):
        p = doc.add_paragraph()
        r1 = p.add_run(f"{label}\t: ")
        r1.bold = True
        r1.font.size = Pt(10)
        r1.font.name = "Calibri"
        for i, line in enumerate(value_lines):
            if i > 0:
                br = p.add_run()
                br.add_break()
            r2 = p.add_run(line)
            r2.font.size = Pt(10)
            r2.font.name = "Calibri"
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)
        return p

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(doc_title)
    run.bold = True
    run.underline = True
    run.font.size = Pt(12)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    doc.add_paragraph()

    if residente_nombre:
        lines = [residente_nombre]
        if residente_cargo:
            lines.append(residente_cargo)
        add_field("A", lines)

    if responsable_nombre:
        lines = [responsable_nombre]
        if responsable_cargo:
            lines.append(responsable_cargo)
        add_field("DE", lines)

    if asunto:
        add_field("ASUNTO", [asunto])

    add_field("FECHA", [f"Santo Domingo de Acobamba, {fecha_label}"])

    doc.add_paragraph()
    add_heading("DATOS GENERALES", level=2, size=10)
    add_field("NOMBRE DEL PROYECTO", [proyecto_nombre])
    add_field("CUI", [cui])
    if clima:
        add_field("CLIMA", [clima])

    if observaciones:
        add_field("OBSERVACIONES", [observaciones])

    for section_data in sections_data:
        s_type = section_data.get("type")

        if s_type == "partidas":
            doc.add_paragraph()
            add_heading(section_data["title"], level=2, size=10)
            avances = section_data["data"]
            if avances:
                hdrs = ["Sector", "Código", "Partida", "Ejecutado por", "Operarios", "Oficiales", "Peones", "Horas", "Cant. Ejec."]
                if show_fecha_col:
                    hdrs = ["Fecha"] + hdrs
                table = doc.add_table(rows=1, cols=len(hdrs))
                table.style = "Light Shading Accent 1"
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                for i, h in enumerate(hdrs):
                    cell = table.rows[0].cells[i]
                    cell.text = h
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.bold = True
                            r.font.size = Pt(8)
                            r.font.name = "Calibri"
                            r.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

                if show_fecha_col:
                    col_widths = [1.5, 1.8, 1.2, 3.8, 1.8, 1.0, 1.0, 1.0, 1.0, 1.0]
                else:
                    col_widths = [1.8, 1.2, 4.5, 1.8, 1.1, 1.1, 1.1, 1.1, 1.1]

                for a in avances:
                    row = table.add_row()
                    vals = []
                    if show_fecha_col:
                        vals.append(a.get("fecha", ""))
                    vals += [
                        a.get("sector_nombre") or "",
                        a.get("partida_codigo") or "",
                        a["partida_nombre"],
                        a.get("subcontrata_nombre") or "Empresa",
                        str(a.get("num_operarios", 0)),
                        str(a.get("num_oficiales", 0)),
                        str(a.get("num_peones", 0)),
                        f'{a.get("horas_trabajadas", 0):.1f}',
                        f'{a.get("cantidad_ejecutada", 0):.2f}',
                    ]
                    for i, v in enumerate(vals):
                        row.cells[i].text = v
                        for p in row.cells[i].paragraphs:
                            for r in p.runs:
                                r.font.size = Pt(8)
                                r.font.name = "Calibri"
                                r.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

                _set_col_widths(table, col_widths)
                doc.add_paragraph()

        elif s_type == "materiales":
            doc.add_paragraph()
            add_heading(section_data["title"], level=2, size=10)
            materiales = section_data["data"]
            if materiales:
                table = doc.add_table(rows=1, cols=4)
                table.style = "Light Shading Accent 1"
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                hdrs = ["Material", "Cantidad", "Unidad", "Fecha"]
                for i, h in enumerate(hdrs):
                    cell = table.rows[0].cells[i]
                    cell.text = h
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.bold = True
                            r.font.size = Pt(8)
                            r.font.name = "Calibri"
                            r.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                for mat in materiales:
                    row = table.add_row()
                    vals = [mat["material"], str(mat["cantidad"]), mat["unidad"], mat["fecha"]]
                    for i, v in enumerate(vals):
                        row.cells[i].text = v
                        for p in row.cells[i].paragraphs:
                            for r in p.runs:
                                r.font.size = Pt(8)
                                r.font.name = "Calibri"
                                r.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                doc.add_paragraph()

        elif s_type == "notas":
            doc.add_paragraph()
            add_heading(section_data["title"], level=2, size=10)
            for nota in section_data["data"]:
                add_para(f"- {nota['nota']}", size=10)

        elif s_type == "fotos":
            doc.add_paragraph()
            add_heading(section_data["title"], level=2, size=10)
            for idx, foto in enumerate(section_data["data"], start=1):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(f"FOTOGRAFIA N° {idx}")
                r.bold = True
                r.font.size = Pt(10)
                r.font.name = "Calibri"
                if os.path.exists(foto["ruta"]):
                    try:
                        doc.add_picture(foto["ruta"], width=Inches(4.5))
                        last_paragraph = doc.paragraphs[-1]
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception:
                        p2 = doc.add_paragraph()
                        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        r2 = p2.add_run("[No se pudo insertar la imagen]")
                        r2.font.size = Pt(9)
                if foto.get("descripcion"):
                    p3 = doc.add_paragraph()
                    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r3 = p3.add_run(f"Descripcion: {foto['descripcion']}")
                    r3.font.size = Pt(10)
                    r3.font.name = "Calibri"

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("_________________________________")
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(responsable_nombre)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(responsable_cargo)
    run.font.size = Pt(10)
    run.font.name = "Calibri"

    return doc


def generar_informe(fecha, avances, materiales, notas, fotos, registro,
                    responsable_nombre, responsable_cargo,
                    residente_nombre=None, residente_cargo=None,
                    asunto="", proyecto_nombre="", cui=""):
    fecha_label = _fmt_fecha_larga(fecha)
    sections = []
    if avances:
        sections.append({"type": "partidas", "title": "AVANCE DE PARTIDAS", "data": avances})
    if materiales:
        sections.append({"type": "materiales", "title": "CONTROL DE MATERIALES", "data": materiales})
    if notas:
        sections.append({"type": "notas", "title": "NOTAS DE CAMPO", "data": notas})
    if fotos:
        sections.append({"type": "fotos", "title": "REGISTRO FOTOGRAFICO", "data": fotos})

    clima = registro.get("clima", "") if registro else ""
    observaciones = registro.get("observaciones", "") if registro else ""

    return _build_doc("REPORTE DIARIO DE OBRA", fecha_label, sections,
                      residente_nombre, residente_cargo,
                      responsable_nombre, responsable_cargo,
                      asunto, proyecto_nombre, cui, clima, observaciones,
                      show_fecha_col=False)


def generar_informe_semanal(fecha_ini, fecha_fin, registros, avances, materiales, notas,
                            responsable_nombre, responsable_cargo,
                            residente_nombre=None, residente_cargo=None,
                            asunto="", proyecto_nombre="", cui=""):
    fecha_label = f"Del {_fmt_fecha(fecha_ini)} al {_fmt_fecha(fecha_fin)}"
    sections = []

    clima = "Varios (ver detalle)"
    observaciones = ""

    if registros:
        lineas = []
        for r in registros:
            lineas.append(f"Dia {_fmt_fecha(r['fecha'])}: Clima {r['clima']}")
        observaciones = f"Dias registrados: {len(registros)}\n" + "\n".join(lineas)

    if avances:
        sections.append({"type": "partidas", "title": "AVANCE DE PARTIDAS", "data": avances})
    if materiales:
        sections.append({"type": "materiales", "title": "CONTROL DE MATERIALES", "data": materiales})
    if notas:
        sections.append({"type": "notas", "title": "NOTAS DE CAMPO", "data": notas})

    return _build_doc("REPORTE SEMANAL DE OBRA", fecha_label, sections,
                      residente_nombre, residente_cargo,
                      responsable_nombre, responsable_cargo,
                      asunto, proyecto_nombre, cui, clima, observaciones,
                      show_fecha_col=True)
